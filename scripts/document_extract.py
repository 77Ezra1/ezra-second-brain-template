#!/usr/bin/env python
from __future__ import annotations

import argparse
import io
import json
import re
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Any

TEXT_SUFFIXES = {".md", ".txt", ".markdown"}
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}
DOCUMENT_SUFFIXES = TEXT_SUFFIXES | PDF_SUFFIXES | DOCX_SUFFIXES | {".pptx", ".xlsx", ".html", ".htm", ".epub"}


def word_count(text: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]", text)) + len(re.findall(r"[A-Za-z0-9_]+", text))


def first_heading(text: str, fallback: str) -> str:
    for line in text.splitlines():
        m = re.match(r"^#{1,6}\s+(.+?)\s*$", line.strip())
        if m:
            return m.group(1).strip()
    return fallback


def maybe_install(package: str) -> None:
    env = dict(**__import__("os").environ)
    # Use user's local proxy convention if present/needed; harmless when proxy is unavailable for cached installs.
    env.setdefault("HTTP_PROXY", "http://127.0.0.1:7897")
    env.setdefault("HTTPS_PROXY", "http://127.0.0.1:7897")
    env.setdefault("ALL_PROXY", "socks5://127.0.0.1:7897")
    subprocess.run(["uv", "pip", "install", package], text=True, capture_output=True, check=True, env=env)


def extract_with_markitdown(path: Path) -> str:
    try:
        from markitdown import MarkItDown
    except ModuleNotFoundError:
        maybe_install("markitdown[all]")
        from markitdown import MarkItDown
    result = MarkItDown().convert(str(path))
    return result.text_content or ""


def extract_with_docling(path: Path) -> str:
    try:
        from docling.document_converter import DocumentConverter
    except ModuleNotFoundError:
        maybe_install("docling")
        from docling.document_converter import DocumentConverter
    result = DocumentConverter().convert(str(path))
    doc = result.document
    if hasattr(doc, "export_to_markdown"):
        return doc.export_to_markdown()
    return str(doc)


def extract_with_marker(path: Path) -> str:
    try:
        from marker.converters.pdf import PdfConverter
        from marker.models import create_model_dict
        from marker.output import text_from_rendered
    except ModuleNotFoundError:
        maybe_install("marker-pdf")
        from marker.converters.pdf import PdfConverter
        from marker.models import create_model_dict
        from marker.output import text_from_rendered
    converter = PdfConverter(artifact_dict=create_model_dict())
    rendered = converter(str(path))
    text, _, _ = text_from_rendered(rendered)
    return text or ""


def extract_with_pymupdf(path: Path) -> str:
    try:
        import pymupdf4llm
        return pymupdf4llm.to_markdown(str(path)) or ""
    except ModuleNotFoundError:
        try:
            maybe_install("pymupdf pymupdf4llm")
            import pymupdf4llm
            return pymupdf4llm.to_markdown(str(path)) or ""
        except Exception:
            try:
                import pymupdf
            except ModuleNotFoundError:
                maybe_install("pymupdf")
                import pymupdf
            doc = pymupdf.open(str(path))
            return "\n\n".join(page.get_text("text") for page in doc)


def extract_with_python_docx(path: Path) -> str:
    try:
        import docx
    except ModuleNotFoundError:
        maybe_install("python-docx")
        import docx
    document = docx.Document(str(path))
    lines: list[str] = []
    for p in document.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = (p.style.name or "").lower() if p.style else ""
        if "heading 1" in style:
            lines.append(f"# {text}")
        elif "heading 2" in style:
            lines.append(f"## {text}")
        else:
            lines.append(text)
    for table in document.tables:
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        if rows:
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
    return "\n\n".join(lines)


def ocr_image_bytes(image_bytes: bytes) -> str:
    try:
        from rapidocr import RapidOCR
    except ModuleNotFoundError:
        maybe_install("rapidocr onnxruntime pillow")
        from rapidocr import RapidOCR
    try:
        from PIL import Image
    except ModuleNotFoundError:
        maybe_install("pillow")
        from PIL import Image
    import numpy as np
    image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    result = RapidOCR()(np.array(image))
    if result is None:
        return ""
    if hasattr(result, "txts"):
        return "\n".join(str(t).strip() for t in result.txts if str(t).strip())
    if isinstance(result, tuple) and result:
        rows = result[0] or []
    else:
        rows = result or []
    texts: list[str] = []
    for row in rows:
        try:
            text = row[1]
        except Exception:
            text = ""
        if text:
            texts.append(str(text).strip())
    return "\n".join(t for t in texts if t)


def ocr_pdf_images(path: Path, max_pages: int = 20) -> list[dict[str, Any]]:
    try:
        import pymupdf
    except ModuleNotFoundError:
        maybe_install("pymupdf")
        import pymupdf
    doc = pymupdf.open(str(path))
    pages: list[dict[str, Any]] = []
    for page_index, page in enumerate(doc[:max_pages], start=1):
        texts: list[str] = []
        images = page.get_images(full=True)
        for img in images:
            xref = img[0]
            try:
                extracted = doc.extract_image(xref)
                text = ocr_image_bytes(extracted.get("image", b""))
                if text:
                    texts.append(text)
            except Exception:
                continue
        if texts:
            pages.append({"page": page_index, "text": "\n".join(texts), "image_count": len(images)})
    return pages


def append_image_ocr(text: str, path: Path, enabled: bool, max_pages: int = 20) -> tuple[str, dict[str, Any] | None]:
    if not enabled or path.suffix.lower() not in PDF_SUFFIXES:
        return text, None
    pages = ocr_pdf_images(path, max_pages=max_pages)
    if not pages:
        return text, {"enabled": True, "pages": 0, "chars": 0}
    block = ["", "## Image OCR", ""]
    for item in pages:
        block.append(f"### Page {item['page']} images")
        block.append(item["text"])
        block.append("")
    ocr_text = "\n".join(block).strip()
    return text.rstrip() + "\n\n" + ocr_text + "\n", {"enabled": True, "pages": len(pages), "chars": len(ocr_text)}


def preferred_methods(path: Path, preferred: list[str] | None = None) -> list[str]:
    if preferred:
        return preferred
    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return ["text"]
    if suffix in PDF_SUFFIXES:
        return ["markitdown", "docling", "marker", "pymupdf"]
    if suffix in DOCX_SUFFIXES:
        return ["markitdown", "docling", "python-docx"]
    return ["markitdown", "docling"]


def extract_by_method(method: str, path: Path) -> str:
    if method == "text":
        return path.read_text(encoding="utf-8", errors="replace")
    if method == "markitdown":
        return extract_with_markitdown(path)
    if method == "docling":
        return extract_with_docling(path)
    if method == "marker":
        return extract_with_marker(path)
    if method == "pymupdf":
        return extract_with_pymupdf(path)
    if method == "python-docx":
        return extract_with_python_docx(path)
    raise ValueError(f"Unknown extraction method: {method}")


def extract_document(
    path: str | Path,
    output: str | Path | None = None,
    preferred: list[str] | None = None,
    min_chars: int = 10,
    ocr_images: bool = False,
    ocr_max_pages: int = 20,
) -> dict[str, Any]:
    src = Path(path).expanduser().resolve()
    if not src.exists():
        return {"ok": False, "path": str(src), "error": "FILE_NOT_FOUND", "attempts": []}
    output_path = Path(output).expanduser().resolve() if output else Path(tempfile.gettempdir()) / f"hermes-document-extract-{src.stem}.md"
    attempts: list[dict[str, Any]] = []
    errors: list[str] = []
    for method in preferred_methods(src, preferred):
        try:
            text = extract_by_method(method, src).strip()
            image_ocr: dict[str, Any] | None = None
            text, image_ocr = append_image_ocr(text, src, ocr_images, max_pages=ocr_max_pages)
            attempts.append({"method": method, "ok": True, "chars": len(text), "image_ocr": image_ocr})
            if len(text) < min_chars:
                errors.append(f"{method}: extracted content too short ({len(text)} chars)")
                continue
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(text, encoding="utf-8")
            result = {
                "ok": True,
                "path": str(src),
                "title": first_heading(text, src.stem),
                "method": method,
                "content_file": str(output_path),
                "word_count": word_count(text),
                "attempts": attempts,
            }
            if image_ocr is not None:
                result["image_ocr"] = image_ocr
            return result
        except Exception as exc:  # noqa: BLE001 - collect fallback reasons
            attempts.append({"method": method, "ok": False, "error": str(exc)})
            errors.append(f"{method}: {exc}")
    return {"ok": False, "path": str(src), "error": "; ".join(errors) or "EXTRACTION_FAILED", "attempts": attempts}


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Extract local documents to Markdown using V2 fallback chain: MarkItDown → Docling/Marker/PyMuPDF/python-docx.")
    parser.add_argument("--file", required=True)
    parser.add_argument("--output")
    parser.add_argument("--preferred", help="Comma-separated extraction methods, e.g. markitdown,docling,pymupdf")
    parser.add_argument("--ocr-images", action="store_true", help="OCR embedded PDF images and append recognized text to Markdown")
    parser.add_argument("--ocr-max-pages", type=int, default=20, help="Max PDF pages to scan for image OCR")
    args = parser.parse_args(argv)
    preferred = [p.strip() for p in args.preferred.split(",") if p.strip()] if args.preferred else None
    result = extract_document(args.file, output=args.output, preferred=preferred, ocr_images=args.ocr_images, ocr_max_pages=args.ocr_max_pages)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if result.get("ok") else 1


if __name__ == "__main__":
    raise SystemExit(main())
